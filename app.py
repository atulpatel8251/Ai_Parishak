import json

from langchain_community.llms import OpenAI
import chat
import openai
from langchain_community.llms import OpenAI
import streamlit as st
import os
import pandas as pd
from rag import *
from PIL import Image
from chat import load_chain
import numpy as np
import logging
from prompts import initialise_prompt
import pytesseract
from pdf2image import convert_from_bytes
from PyPDF2 import PdfReader
from langchain_openai import ChatOpenAI
from langchain.chains import ConversationChain
from langchain.memory import ConversationBufferMemory
import docx
from prompts import ai_prompt,ai_topic_prompt,latex_prompt,student_prompt,mcq_test_prompt,learn_outcome_prompt,ai_topic_prompt1,ai_topic_prompt2,ai_topic_prompt3,ai_topic_prompt5
from functions import read_word_table,create_word_doc,get_text
import subprocess
import re
import constants
from dotenv import load_dotenv
from docxlatex import Document
import json
import base64
import fixed_function
from io import BytesIO
import streamlit as st
from PyPDF2 import PdfReader
from langchain_community.embeddings import OpenAIEmbeddings
from langchain.text_splitter import CharacterTextSplitter
from langchain.chains.question_answering import load_qa_chain
from langchain.llms import OpenAI
from langchain_community.vectorstores import FAISS
from docx import Document
from langchain.vectorstores import Chroma
from translate import Translator
import os
import time
from io import StringIO
import shutil
#from docxlatex import Document
# from PIL import Image
#client = OpenAI()

st.set_page_config(layout='wide')

hide_streamlit_style = """
            <style>
            
            footer {visibility: hidden;}
            h1 {
               color: green;
               font-size: 35px;    
               }
            
            </style>
            """
st.markdown(hide_streamlit_style, unsafe_allow_html=True)
#add_selectbox = st.sidebar.selectbox(
#    "Store the chat hsitory",
#    ("Teachers", "Students", "Administration")
#)

load_dotenv()

images = ['6MarkQ']


#os.environ["OPENAI_API_TYPE"] = ""
#os.environ["OPENAI_API_VERSION"] = ""openai_api_key = os.getenv("OPENAI_API_KEY")
openai_api_key2 = st.secrets["secret_section"]["OPENAI_API_KEY"]

#openai_api_key = os.getenv('OPENAI_API_KEY')
TEMP_MD_FILE = r"question.pdf"
TEMP_PDF_FILE = "__temp.pdf"

def get_base64(bin_file):
    with open(bin_file, 'rb') as f:
        data = f.read()
    return  base64.b64encode(data).decode()



def download_doc(doc):
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer


with open('style.css') as f:
 st.markdown(f'<style>{f.read()}</style>', unsafe_allow_html=True) 


def mcq_to_markdown(mcq_dict):
    markdown = ""
    for key, value in mcq_dict.items():
        if key.isdigit():
            markdown += f"Question {key}: {value}\n"
        elif key.startswith('image'):
            markdown += f"![Image {key[-1]}]({value})\n\n"
    return markdown


def display_mcq(mcq_dict):
    for key, value in mcq_dict.items():
        st.markdown(f"""**Question {key}:** {value['question']}""")
        c1,c2 = st.columns([3,1])
        with c1:
            st.write("Options:")
            for option_key, option_value in value['options'].items():
                st.write(f"{option_key}. {option_value}")
        with c2:
            st.image(f"{value['image']}", use_column_width=True)

def is_word_in_text(word, text):
    """
    Check if a word is within a given text.

    Args:
    - word (str): The word to check for.
    - text (str): The text to search within.

    Returns:
    - bool: True if the word is found in the text, False otherwise.
    """
    # Split the text into individual words
    words_in_text = text.split()

    # Check if the word is in the list of words from the text
    if word in words_in_text:
        return True
    else:
        return False
    
def text_to_latex(question):
    # Convert question to LaTeX format
    latex_question = r"\text{" + question + r"}"
    return f"{latex_question}"
    
def add_dollar_signs(text):
    # Regular expression pattern to find equations
    equation_pattern = r"([+-]?\s*\d*\s*[xX]\^\d+)"

    # Find equations in the text using regular expression
    equations = re.findall(equation_pattern, text)

    # Replace each equation with the same equation surrounded by dollar signs
    for equation in equations:
        text = text.replace(equation, f"${equation}$")

    return text
    
def generate_quiz(text):

    test_source_text = text
    quiz_rag = RAGTask(task_prompt_builder=revision_quiz_json_builder)
#     summary_rag = RAGTask(task_prompt_builder=plaintext_summary_builder)
#     glossary_rag = RAGTask(task_prompt_builder=get_glossary_builder)

    outputs = []
    # for rag_task in [quiz_rag, summary_rag, glossary_rag]:
    #     output = rag_task.get_output(source_text=test_source_text)
    #     outputs.append(output)
    for rag_task in [quiz_rag]:
        output = rag_task.get_output(source_text=test_source_text)
        outputs.append(output)
    return outputs

def decrement_question_num():
    if st.session_state['curr_question'] > 0:
        st.session_state['curr_question'] -= 1
        #ClearAll()

def increment_question_num():
    print('Incrementing question', st.session_state['curr_question'])
    if st.session_state['curr_question'] < st.session_state['quiz_length'] - 1:
        st.session_state['curr_question'] += 1
        #ClearAll()
    


def display_q(df):
    for index,row in df.iterrows():
        st.markdown("##### **Question "+str(index+1)+"**")
        if row['Image Path']=='no_image':
            # st.session_state.format_chain = ConversationChain( llm=AzureChatOpenAI(
            # deployment_name="gpt-4turbo",
            # temperature=0
            # ))
            # formatted_question = chat.format_response(row['Question'])
            # st.latex(formatted_question)
            st.info(row['Question'])
        if row['Image Path'] !='no_image':
            # st.session_state.format_chain = ConversationChain( llm=AzureChatOpenAI(
            # deployment_name="gpt-4turbo",
            # temperature=0
            # ))
            # formatted_question = chat.format_response(row['Question'])
            # st.latex(formatted_question)
            st.info(row['Question'])
            st.image(row['Image Path'],width=180)
        st.divider()

pd.set_option('display.max_colwidth', None)
st.session_state.ques = pd.read_csv('update_question.csv')



with st.container():
    col1, col2, col3 = st.columns([0.2, 0.6, 0.2], gap="small")
    with col1:
        logo_image = Image.open('assests/madhya-pradesh-logo.png')
        resized_logo = logo_image.resize((150, 150))
        st.image(resized_logo)
    with col2:
        st.markdown("# GyanKosh")
        st.markdown("###### AI Based Question Generation Assistance")
    with col3:
        l = Image.open('assests/28072020125926mpsedclogo.png')
        re = l.resize((165, 127))  # Corrected the resize method call
        st.image(re)
         
        

def on_text_area_change():
    st.session_state.page_text = st.session_state.my_text_area

def read_pdf_page(file, page_number):
    pdfReader = PdfReader(file)
    page = pdfReader.pages[page_number]
    return page.extract_text()




def markdown_to_pdf(markdown: str, output_file: str):
    """
    Convert Markdown to PDF
    :param markdown: Markdown string
    :param output_file: Output file
    """
    with open(TEMP_MD_FILE, "w",encoding='utf-8') as f:
        f.write(markdown)


def correct_bhashni_translations(text,lowercase_dict):
    corrected_text = []
    words = text.split()  # Tokenize the input text into words
    for word in words:
        # Check if the word needs correction
        if word in lowercase_dict:
            corrected_text.append(lowercase_dict[word])
        else:
            corrected_text.append(word)  # If no correction needed, keep the word unchanged
    return ' '.join(corrected_text)

def split_text_into_chunks(text, max_length):
    words = text.split()
    chunks = []
    current_chunk = []
    current_length = 0
    for word in words:
        if current_length + len(word) + 1 <= max_length:
            current_chunk.append(word)
            current_length += len(word) + 1
        else:
            chunks.append(" ".join(current_chunk))
            current_chunk = [word]
            current_length = len(word) + 1
    if current_chunk:
        chunks.append(" ".join(current_chunk))
    return chunks
from PyPDF2 import PdfReader
def list_files(folder_path):
    return os.listdir(folder_path)

# Function to remove file extension
def remove_extension(filename):
    return os.path.splitext(filename)[0]

# def setup_tesseract():
#     """Set up Tesseract environment variables and check installation"""
#     # Modify these paths according to your system
#     tesseract_path = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
#     tessdata_path = r"C:\Program Files\Tesseract-OCR\tessdata"
    
#     if not os.path.exists(tesseract_path):
#         raise EnvironmentError("Tesseract executable not found. Please install Tesseract OCR.")
    
#     if not os.path.exists(tessdata_path):
#         raise EnvironmentError("Tessdata directory not found. Please install language files.")
    
#     # Set environment variables
#     os.environ['TESSDATA_PREFIX'] = tessdata_path
#     pytesseract.pytesseract.tesseract_cmd = tesseract_path

# def check_language_files(languages):
#     """Check if required language files exist"""
#     tessdata_path = os.environ.get('TESSDATA_PREFIX')
#     missing_langs = []
    
#     for lang in languages:
#         lang_file = os.path.join(tessdata_path, f"{lang}.traineddata")
#         if not os.path.exists(lang_file):
#             missing_langs.append(lang)
    
#     return missing_langs
from pdf2image import convert_from_path

# def extract_text_with_ocr(pdf_file_path):
#     """Extract text from PDF using OCR with proper error handling"""
#     try:
#         # Setup Tesseract environment
#         setup_tesseract()
        
#         # Check for required language files
#         required_langs = ['hin', 'eng']
#         missing_langs = check_language_files(required_langs)
        
#         if missing_langs:
#             raise EnvironmentError(
#                 f"Missing language files for: {', '.join(missing_langs)}. "
#                 "Please download them from https://github.com/tesseract-ocr/tessdata"
#             )
        
#         # Convert PDF to images
#         images = convert_from_path(pdf_file_path)
#         all_text = []
        
#         # Process each page
#         for i, img in enumerate(images, 1):
#             try:
#                 text = pytesseract.image_to_string(img, lang='hin+eng')
#                 all_text.append(text)
#                 #st.success(f"Successfully processed page {i}")
#             except Exception as e:
#                 st.warning(f"Error processing page {i}: {str(e)}")
        
#         return "\n".join(all_text)
    
#     except Exception as e:
#         error_msg = str(e)
#         st.error(f"Error during OCR extraction: {error_msg}")
        
#         # Provide more specific guidance based on the error
#         if "TESSDATA_PREFIX" in error_msg:
#             st.info("Please follow these steps to fix the issue:\n"
#                    "1. Download language files from https://github.com/tesseract-ocr/tessdata\n"
#                    "2. Place them in the Tesseract tessdata directory\n"
#                    "3. Set the TESSDATA_PREFIX environment variable")
#         elif "tesseract.exe" in error_msg.lower():
#             st.info("Tesseract OCR needs to be installed. Visit: https://github.com/UB-Mannheim/tesseract/wiki")
            
#         return ""
import hashlib
from datetime import datetime
import multiprocessing
from concurrent.futures import ThreadPoolExecutor
import concurrent.futures


class OCRCache:
    def __init__(self, cache_dir="./ocr_cache"):
        self.cache_dir = cache_dir
        self.cache_index_file = os.path.join(cache_dir, "cache_index.json")
        self.initialize_cache()
    
    def initialize_cache(self):
        # Initialize session state cache
        if 'ocr_cache' not in st.session_state:
            st.session_state.ocr_cache = {}
        
        # Initialize file cache
        os.makedirs(self.cache_dir, exist_ok=True)
        if not os.path.exists(self.cache_index_file):
            self.save_cache_index({})
    
    def get_file_hash(self, file_path):
        """Generate hash of file content and modification time"""
        modification_time = os.path.getmtime(file_path)
        file_size = os.path.getsize(file_path)
        with open(file_path, 'rb') as f:
            file_content = f.read(1024 * 1024)  # Read first MB for hashing
        hash_string = f"{file_path}_{modification_time}_{file_size}_{file_content}"
        return hashlib.md5(hash_string.encode()).hexdigest()
    
    def is_cache_valid(self, file_path, file_hash):
        """Check if cache is valid for given file"""
        cache_index = self.load_cache_index()
        if file_hash in cache_index:
            cache_file = os.path.join(self.cache_dir, f"{file_hash}.txt")
            if os.path.exists(cache_file):
                cached_info = cache_index[file_hash]
                # Check if file path and modification time match
                if (cached_info['file_path'] == file_path and 
                    os.path.getmtime(file_path) == cached_info.get('modification_time')):
                    return True
        return False
    
    def load_cache_index(self):
        """Load cache index from file"""
        try:
            with open(self.cache_index_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception:
            return {}
    
    def save_cache_index(self, index):
        """Save cache index to file"""
        try:
            with open(self.cache_index_file, 'w', encoding='utf-8') as f:
                json.dump(index, f, ensure_ascii=False, indent=2)
        except Exception as e:
            st.warning(f"Could not save cache index: {str(e)}")
    
    def get_cached_text(self, file_path):
        """Retrieve cached text from memory or file"""
        try:
            file_hash = self.get_file_hash(file_path)
            
            # First check if cache is valid
            if not self.is_cache_valid(file_path, file_hash):
                return None
            
            # Try memory cache first
            if file_hash in st.session_state.ocr_cache:
                st.success(f"Using memory cache for {os.path.basename(file_path)}")
                return st.session_state.ocr_cache[file_hash]
            
            # Try file cache
            cache_file = os.path.join(self.cache_dir, f"{file_hash}.txt")
            if os.path.exists(cache_file):
                try:
                    with open(cache_file, 'r', encoding='utf-8') as f:
                        text = f.read()
                        # Store in memory cache for future use
                        st.session_state.ocr_cache[file_hash] = text
                        st.success(f"Using file cache for {os.path.basename(file_path)}")
                        return text
                except Exception as e:
                    st.warning(f"Error reading cache file: {str(e)}")
                    return None
            return None
        except Exception as e:
            st.warning(f"Cache retrieval error: {str(e)}")
            return None

    def save_text_to_cache(self, file_path, text):
        """Save extracted text to both memory and file cache"""
        try:
            file_hash = self.get_file_hash(file_path)
            
            # Check if already cached and valid
            if self.is_cache_valid(file_path, file_hash):
                return  # Skip if already properly cached
            
            # Save to memory cache
            st.session_state.ocr_cache[file_hash] = text
            
            # Save to file cache
            cache_file = os.path.join(self.cache_dir, f"{file_hash}.txt")
            with open(cache_file, 'w', encoding='utf-8') as f:
                f.write(text)
            
            # Update cache index
            cache_index = self.load_cache_index()
            cache_index[file_hash] = {
                'file_path': file_path,
                'modification_time': os.path.getmtime(file_path),
                'cached_date': datetime.now().isoformat(),
                'cache_file': f"{file_hash}.txt"
            }
            self.save_cache_index(cache_index)
            st.success(f"Cached text for {os.path.basename(file_path)}")
            
        except Exception as e:
            st.warning(f"Could not save to cache: {str(e)}")

def extract_text_with_ocr_cached(pdf_file_path, cache_system):
    """Extract text from PDF using hybrid cache system"""
    # Check cache first
    cached_text = cache_system.get_cached_text(pdf_file_path)
    if cached_text is not None:
        st.info(f"Using cached text for {os.path.basename(pdf_file_path)}")
        return cached_text
    
    # If not in cache, perform OCR
    try:
        images = convert_from_path(
            pdf_file_path,
            dpi=150,
            thread_count=2,
            grayscale=True,
            size=(1600, None)
        )
        
        max_workers = min(2, len(images))
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            results = list(executor.map(process_page, images))
        
        extracted_text = "\n".join(filter(None, results))
        
        # Save to cache if extraction was successful
        if extracted_text.strip():
            cache_system.save_text_to_cache(pdf_file_path, extracted_text)
        
        return extracted_text
    
    except Exception as e:
        st.error(f"Error during OCR extraction: {str(e)}")
        return ""
    
def optimize_image_for_ocr(image):
    """Optimize image for faster OCR processing"""
    # Convert to grayscale if not already
    if image.mode != 'L':
        image = image.convert('L')
    
    # Resize image if too large (maintain aspect ratio)
    max_dimension = 2000
    if max(image.size) > max_dimension:
        ratio = max_dimension / max(image.size)
        new_size = tuple(int(dim * ratio) for dim in image.size)
        image = image.resize(new_size, Image.LANCZOS)
    
    # Improve contrast
    image = Image.fromarray(np.uint8(np.clip((np.array(image) * 1.2), 0, 255)))
    
    return image

import os
import pathlib
from PIL import Image
import pytesseract
import streamlit as st

def setup_tesseract(base_path="./Tesseract-OCR"):
    """
    Configure Tesseract for both local and cloud environments
    """
    try:
        # Check for Streamlit Cloud's Tesseract installation
        if os.path.exists('/usr/bin/tesseract'):
            pytesseract.pytesseract.tesseract_cmd = '/usr/bin/tesseract'
        else:
            # For local development
            pytesseract.pytesseract.tesseract_cmd = 'tesseract'
        
        # Quick test
        test_image = Image.new('RGB', (1, 1), color='white')
        pytesseract.image_to_string(test_image)
        st.success("Tesseract setup completed successfully!")
        return True
    except Exception as e:
        st.error(f"Tesseract setup failed: {str(e)}")
        return False
                
    except Exception as e:
        st.error(f"""Tesseract setup failed. Please check:
        1. Tesseract is installed in: {base_path}
        2. Language files are present in: {tessdata_dir}
        
        Error: {str(e)}""")
        return False
    

def process_page(img, language='hin+eng'):
    """Process a single page with cloud-optimized settings"""
    try:
        # Optimize image
        img = optimize_image_for_ocr(img)
        
        # OCR with optimized settings
        custom_config = r'--oem 3 --psm 6 -c preserve_interword_spaces=1'
        try:
            text = pytesseract.image_to_string(img, lang=language, config=custom_config)
        except Exception:
            st.warning(f"Failed with language {language}, falling back to English")
            text = pytesseract.image_to_string(img, lang='eng', config=custom_config)
        
        return text.strip()
    except Exception as e:
        st.error(f"Error processing page: {str(e)}")
        return ""

def optimize_image_for_ocr(image):
    """Optimize image for cloud processing"""
    # Convert to grayscale if not already
    if image.mode != 'L':
        image = image.convert('L')
    
    # Resize if too large (maintain aspect ratio)
    max_dimension = 1600  # Reduced for cloud processing
    if max(image.size) > max_dimension:
        ratio = max_dimension / max(image.size)
        new_size = tuple(int(dim * ratio) for dim in image.size)
        image = image.resize(new_size, Image.LANCZOS)
    
    # Improve contrast
    image = Image.fromarray(np.uint8(np.clip((np.array(image) * 1.2), 0, 255)))
    
    return image

def extract_text_with_ocr_optimized(pdf_file_path):
    """Extract text from PDF using cache if available"""
    # Check cache first
    cached_text = cache_system.get_cached_text(pdf_file_path)
    if cached_text is not None:
        st.info(f"Using cached text for {os.path.basename(pdf_file_path)}")
        return cached_text
    
    # If not in cache, perform OCR
    try:
        # Convert PDF to images with cloud-optimized settings
        images = convert_from_path(
            pdf_file_path,
            dpi=150,  # Reduced DPI for cloud processing
            thread_count=2,  # Limited threads for cloud environment
            grayscale=True,
            size=(1600, None)  # Reduced size for cloud processing
        )
        
        # Process with limited workers for cloud environment
        max_workers = min(2, len(images))
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            results = list(executor.map(process_page, images))
        
        extracted_text = "\n".join(filter(None, results))
        
        # Save to cache if extraction was successful
        if extracted_text.strip():
            cache_system.save_text_to_cache(pdf_file_path, extracted_text)
        
        return extracted_text
    
    except Exception as e:
        st.error(f"Error during OCR extraction: {str(e)}")
        return ""

def batch_process_pdfs(selected_files, folder_path, progress_bar, status_text):
    """Process multiple PDFs in parallel"""
    total_files = len(selected_files)
    combined_text = []
    processed_files = []
    
    # Process files in smaller batches to manage memory
    batch_size = 3
    for i in range(0, total_files, batch_size):
        batch = selected_files[i:i + batch_size]
        
        # Process batch in parallel
        with ThreadPoolExecutor(max_workers=batch_size) as executor:
            future_to_file = {
                executor.submit(
                    extract_text_with_ocr_optimized, 
                    os.path.join(folder_path, file + '.pdf')
                ): file for file in batch
            }
            
            for future in concurrent.futures.as_completed(future_to_file):
                file = future_to_file[future]
                try:
                    text = future.result()
                    if text.strip():
                        combined_text.append(text)
                        processed_files.append(file)
                    
                    # Update progress
                    progress = (len(processed_files) / total_files)
                    progress_bar.progress(progress)
                    status_text.text(f"Processed {len(processed_files)}/{total_files} files")
                    
                except Exception as e:
                    st.warning(f"Error processing {file}: {str(e)}")
    
    return combined_text, processed_files    

def batch_process_pdfs_with_cache(selected_files, folder_path, progress_bar, status_text):
    """Process multiple PDFs using cache when available"""
    total_files = len(selected_files)
    combined_text = []
    processed_files = []
    
    # Initialize cache system
    cache_system = OCRCache()
    
    # Track which files need processing
    files_to_process = []
    cached_texts = {}
    
    # First check cache for all files
    for file in selected_files:
        file_path = os.path.join(folder_path, file + '.pdf')
        cached_text = cache_system.get_cached_text(file_path)
        
        if cached_text is not None:
            cached_texts[file] = cached_text
        else:
            files_to_process.append(file)
    
    # Update progress for cached files
    cached_count = len(cached_texts)
    if cached_count > 0:
        progress = (cached_count / total_files)
        progress_bar.progress(progress)
        status_text.text(f"Retrieved {cached_count} files from cache")
    
    # Process only uncached files
    if files_to_process:
        batch_size = 3
        for i in range(0, len(files_to_process), batch_size):
            batch = files_to_process[i:i + batch_size]
            
            with ThreadPoolExecutor(max_workers=batch_size) as executor:
                future_to_file = {
                    executor.submit(
                        extract_text_with_ocr_cached,
                        os.path.join(folder_path, file + '.pdf'),
                        cache_system
                    ): file for file in batch
                }
                
                for future in concurrent.futures.as_completed(future_to_file):
                    file = future_to_file[future]
                    try:
                        text = future.result()
                        if text.strip():
                            cached_texts[file] = text
                        
                        # Update progress including cached files
                        progress = (len(cached_texts) / total_files)
                        progress_bar.progress(progress)
                        status_text.text(f"Processed {len(cached_texts)}/{total_files} files")
                        
                    except Exception as e:
                        st.warning(f"Error processing {file}: {str(e)}")
    
    # Combine texts in original order
    for file in selected_files:
        if file in cached_texts:
            combined_text.append(cached_texts[file])
            processed_files.append(file)
    
    return combined_text, processed_files    

# Function to convert PDF to text
def pdf_to_text(file_path):
    text = ""
    with open(file_path, 'rb') as file:
        pdf_reader = PdfReader(file)
        for page in pdf_reader.pages:
            text += page.extract_text()
    return text

# Helper function: Check if embeddings exist in the vector DB
def check_embeddings_exist(vector_store, subject, chapter):
    # Perform a simple similarity search
    results = vector_store.similarity_search("dummy query", k=10)  # Adjust k based on your need

    # Filter the results based on metadata (subject and chapter)
    for result in results:
        if result.metadata.get("subject") == subject and result.metadata.get("chapter") == chapter:
            return True  # Embeddings for the given subject and chapter exist

    return False 

# Helper function: Create embeddings from text and store in Chroma only if not already stored
def create_and_store_embeddings(text, subject, chapter, vector_store):
    # Check if the embeddings already exist for the specific subject and chapter
    if check_embeddings_exist(vector_store, subject, chapter):
        #st.info("file already exist.")
        return  # Exit the function if the file already exists

    # If the embeddings do not exist, create embeddings
    else:
        embedding_model = OpenAIEmbeddings(openai_api_key=openai_api_key2)

        # Split the text into manageable chunks (Chroma works well with smaller chunks)
        text_chunks = [text[i:i + 1000] for i in range(0, len(text), 1000)]  # Adjust chunk size as needed
        
        # Generate embeddings for the document chunks
        embeddings = embedding_model.embed_documents(text_chunks)
        
        # Store the generated embeddings in the vector database with metadata
        vector_store.add_texts(
            texts=text_chunks,
            embeddings=embeddings,
            metadatas=[{"subject": subject, "chapter": chapter} for _ in text_chunks]
        )
        
        #st.info("Embeddings created and stored successfully!")

# Initialize Chroma vector store
def initialize_chroma():
    # Chroma DB configuration (you can use any other vector DB like Pinecone, etc.)
    chroma_vector_store = Chroma(persist_directory="./chroma_db", embedding_function=OpenAIEmbeddings(openai_api_key=openai_api_key2))
    return chroma_vector_store

def create_excel(data, filename="questions.xlsx"):
    # Convert data to a DataFrame with specified columns
    df = pd.DataFrame(data, columns=["Unit", "Level", "Topic", "Question", "Option 1", "Option 2", "Option 3", "Option 4", "Correct Answer"])
    
    # Use BytesIO for in-memory storage of the Excel file
    excel_buffer = BytesIO()
    with pd.ExcelWriter(excel_buffer, engine="xlsxwriter") as writer:
        df.to_excel(writer, index=False, sheet_name="Questions")
    
    excel_buffer.seek(0)  # Rewind buffer for reading
    return excel_buffer

def clean_question_text(text):
    """Clean question text without using regex"""
    text = text.strip()
    # Remove common question number patterns
    prefixes = ['1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.', '10.',
                'Q1.', 'Q2.', 'Q3.', 'Q4.', 'Q5.', 'Q6.', 'Q7.', 'Q8.', 'Q9.', 'Q10.',
                'Question 1:', 'Question 2:', 'Question 3:', 'Question 4:', 'Question 5:']
    
    for prefix in prefixes:
        if text.startswith(prefix):
            text = text[len(prefix):].strip()
            break
            
    return text

def parse_question(text):
    """Helper function to parse a single question text into its components"""
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    
    # Initialize dictionary with None values
    question_dict = {
        "Question": None,
        "Option 1": None,
        "Option 2": None,
        "Option 3": None,
        "Option 4": None,
        "Correct Answer": None
    }
    
    try:
        # Process first line as question
        if lines:
            question_dict["Question"] = clean_question_text(lines[0])
        
        # Process options
        for i, line in enumerate(lines[1:5], 1):
            line = line.strip()
            # Remove option letters/numbers at the start
            if line.startswith(('a)', 'b)', 'c)', 'd)', 'A)', 'B)', 'C)', 'D)')):
                option_text = line[2:].strip()
            elif line.startswith(('a.', 'b.', 'c.', 'd.', 'A.', 'B.', 'C.', 'D.')):
                option_text = line[2:].strip()
            elif line.startswith(('a', 'b', 'c', 'd', 'A', 'B', 'C', 'D')):
                option_text = line[1:].strip()
            else:
                option_text = line
            question_dict[f"Option {i}"] = option_text
        
        # Process answer
        for line in lines[5:]:
            line = line.lower().strip()
            if line.startswith(('answer:', 'correct answer:', 'ans:')):
                # Extract answer letter
                answer_text = line.replace('answer:', '').replace('correct answer:', '').replace('ans:', '').strip()
                if answer_text in ['a', 'b', 'c', 'd']:
                    option_number = ord(answer_text) - ord('a') + 1
                    question_dict["Correct Answer"] = question_dict[f"Option {option_number}"]
                else:
                    question_dict["Correct Answer"] = answer_text.capitalize()
                break
    except Exception as e:
        print(f"Error parsing question: {str(e)}")
        return None
    
    return question_dict

def list_files(folder_path):
    return os.listdir(folder_path)

def remove_extension(filename):
    return os.path.splitext(filename)[0]

# Initialize session state variables
if 'selected_option' not in st.session_state:
    st.session_state.selected_option = None

if 'quesai' not in st.session_state:
    st.session_state.quesai = None

if 'selected_file' not in st.session_state:
    st.session_state.selected_file = "Select document"
   
st.sidebar.header("Select Module")
st.session_state.teach = st.sidebar.selectbox(
    "",
    ('Teachers', 'Students', 'Administration'),
    key='airadio1'
)
if st.session_state.teach == 'Teachers':
    st.session_state.quesai = st.title("Generate Question and Answer")

    if st.session_state.quesai:
        st.session_state.selected_option = st.radio(
            "Select Options",
            ("Topic Based Questions", "Examination" , "Text Analyzer","Pre Uploaded" , "Terminologies", "Learning Outcomes","Question Paper Generator"),
            horizontal=True,
            index=0,
            key='option'
        )

        choose = st.session_state.selected_option

        col_11, col_22 = st.columns([2, 1])

        with col_11:
            if choose == "Pre Uploaded":
                os.environ['OMP_THREAD_LIMIT'] = str(multiprocessing.cpu_count())

                # Define the base folder based on medium (Hindi or English)
                medium_folder = "./preuploaded"
                medium_options = ["Select Medium", "Hindi Medium", "English Medium"]
                selected_medium = st.selectbox("Select Medium", medium_options, index=0, key="medium_selector")

                if selected_medium != "Select Medium":
                    # Adjust the folder path based on the selected medium
                    subjects_folder = os.path.join(medium_folder, selected_medium)
                    subjects_list = [d for d in os.listdir(subjects_folder) if os.path.isdir(os.path.join(subjects_folder, d))]
                    subjects_list.sort()
                    subjects_list.insert(0, "Select subject")

                    selected_subject = st.selectbox("Select a subject", subjects_list, index=0, key="subject_selector")

                    if selected_subject and selected_subject != "Select subject":
                        folder_path = os.path.join(subjects_folder, selected_subject)
                        files_list = list_files(folder_path)
                        files_list = [remove_extension(filename) for filename in files_list]
                        files_list.insert(0, "Select documents")

                        selected_file = st.multiselect(
                            "Select files (or select all)",
                            ["Select All"] + files_list,
                            key="terminologies_selected_files"
                        )

                        if "Select All" in selected_file:
                            selected_files = files_list[1:]
                        elif "Select documents" in selected_file:
                            selected_file.remove("Select documents")

                        vector_store = initialize_chroma()  # Replace with your initialization logic

                        # Initialize combined_text if it doesn't exist
                        if "combined_text" not in st.session_state:
                            st.session_state.combined_text = None
                        
                        if selected_file:
                            progress_bar = st.progress(0)
                            status_text = st.empty()

                            combined_text, processed_files = batch_process_pdfs_with_cache(
                                selected_file,
                                folder_path,
                                progress_bar,
                                status_text
                            )
                            if combined_text:
                                status_text.text("Generating terminologies and keyterms...")
                                st.session_state.final_text = "\n\n".join(combined_text)

                            # Handling form for question generation
                            with st.form(key="Pre Uploaded"):
                                col1, col2 = st.columns(2)
                                with col1:
                                    st.session_state.complexity = st.selectbox('Complexity Mode Required?*', ['Easy', 'Difficult'], index=0, key="mode")
                                    st.session_state.no_of_questions = st.number_input('No. of Questions to generate*', key="ai_questions", step=1, max_value=30)
                                    st.session_state.mode_of_questions = st.selectbox('Choose Answer Required/Not*', ['Only Questions', 'Questions with Answers'], index=0, key="quesansw")
                                with col2:
                                    st.session_state.type_of_questions = st.selectbox('Choose Question Type*', ['Short Questions', 'Long Questions', 'MCQ', 'Fill in the Blanks', 'True and False'], index=0)
                                submitted = st.form_submit_button("Submit")

                            if submitted and st.session_state.final_text and st.session_state.mode_of_questions != 'Select Option':
                                if st.session_state.final_text:
                                    st.session_state.llm = ConversationChain(llm=ChatOpenAI(model="gpt-4o", temperature=0.7, api_key=openai_api_key2))
                                    chapter_info = f"Chapter: {selected_file}" if selected_file != "All Chapters" else "All Chapters"

                                    # Determine language based on medium
                                    language = "Hindi" if selected_medium == "Hindi" else "English"

                                    # Generate the formatted output based on selected options
                                    formatted_output = st.session_state.llm.predict(input=ai_topic_prompt1.format(
                                        chapter_info,
                                        st.session_state.no_of_questions,
                                        st.session_state.final_text,
                                        language,
                                        st.session_state.mode_of_questions,
                                        st.session_state.type_of_questions,
                                        st.session_state.complexity,
                                        st.session_state.no_of_questions
                                    ))

                                    st.info(formatted_output)
                                    markdown_to_pdf(formatted_output, 'question.pdf')
                                    word_doc = create_word_doc(formatted_output)
                                    doc_buffer = download_doc(word_doc)

                                    st.download_button(
                                        label="Download Word Document",
                                        data=doc_buffer,
                                        file_name="generated_document.docx",
                                        mime="application/octet-stream",
                                        key='worddownload'
                                    )
                                else:
                                    st.info("No relevant results found based on the subject and chapter metadata.")

            if choose == "Terminologies":
    # Set up environment variable for better performance
                os.environ['OMP_THREAD_LIMIT'] = str(multiprocessing.cpu_count())

                # Define the base folder based on medium (Hindi or English)
                medium_folder = "./preuploaded"
                medium_options = ["Select Medium", "Hindi Medium", "English Medium"]
                selected_medium = st.selectbox("Select Medium", medium_options, index=0, key="medium_selector")

                if selected_medium != "Select Medium":
                    # Adjust the folder path based on the selected medium
                    language = "Hindi" if selected_medium == "Hindi Medium" else "English"
                    subjects_folder = os.path.join(medium_folder, selected_medium)
                    subjects_list = [d for d in os.listdir(subjects_folder) if os.path.isdir(os.path.join(subjects_folder, d))]
                    subjects_list.sort()
                    subjects_list.insert(0, "Select subject")

                    selected_subject = st.selectbox("Select a subject", subjects_list, index=0, key="subject_selector")

                    if selected_subject and selected_subject != "Select subject":
                        folder_path = os.path.join(subjects_folder, selected_subject)
                        files_list = list_files(folder_path)
                        files_list = [remove_extension(filename) for filename in files_list]
                        files_list.insert(0, "Select documents")

                        selected_file = st.multiselect(
                            "Select files (or select all)",
                            ["Select All"] + files_list,
                            key="terminologies_selected_files"
                        )

                        if "Select All" in selected_file:
                            selected_files = files_list[1:]  # Exclude "Select documents"
                        else:
                            selected_files = [f for f in selected_file if f != "Select documents"]

                        if selected_files:  # Check if any files are selected
                            progress_bar = st.progress(0)
                            status_text = st.empty()

                            # Initialize session state for storing results
                            if 'combined_text' not in st.session_state:
                                st.session_state.combined_text = []
                            
                            # Process PDFs and get combined text
                            combined_text, processed_files = batch_process_pdfs_with_cache(
                                selected_files,
                                folder_path,
                                progress_bar,
                                status_text
                            )

                            if combined_text:
                                status_text.text("Generating terminologies and keyterms...")
                                final_text = "\n\n".join(combined_text)

                                try:
                                    # Initialize OpenAI model with correct model name
                                    llm = ChatOpenAI(
                                        model="gpt-4o",  # Fixed model name from "gpt-4o" to "gpt-4"
                                        temperature=0.7,
                                        api_key=openai_api_key2
                                    )
                                    # Create conversation chain
                                    st.session_state.mcq_chain = ConversationChain(llm=llm)
                                    
                                    # Generate terminologies
                                    formatted_response = st.session_state.mcq_chain.predict(
                                        input=mcq_test_prompt.format(
                                            final_text,
                                            language
                                              # Default to English if not set
                                        )
                                    )
                                    
                                    # Display the terminologies
                                    st.markdown("### Generated Terminologies and Key Terms")
                                    st.markdown(formatted_response)
                                    
                                    # Generate output files
                                    output_filename = "_".join(processed_files) if len(processed_files) <= 3 else f"{processed_files[0]}_and_{len(processed_files)-1}_others"
                                    
                                    # Create downloads
                                    markdown_to_pdf(formatted_response, f'{output_filename}_terminologies.pdf')
                                    word_doc = create_word_doc(formatted_response)
                                    doc_buffer = download_doc(word_doc)
                                    
                                    col1, col2 = st.columns(2)
                                    with col1:
                                        st.download_button(
                                            label="Download Word Document",
                                            data=doc_buffer,
                                            file_name=f"{output_filename}_terminologies.docx",
                                            mime="application/octet-stream",
                                            key='worddownload_combined'
                                        )
                                    
                                    progress_bar.empty()
                                    status_text.empty()
                                    st.success(f"Successfully processed {len(processed_files)} files and generated terminologies")
                                
                                except Exception as e:
                                    st.error(f"Error generating terminologies: {str(e)}")
                                    st.error("Please check your OpenAI API key and model access.")
                            else:
                                st.error("No text could be extracted from any of the selected files")


            if choose == "Learning Outcomes":
    # Set up environment variable for better performance
                os.environ['OMP_THREAD_LIMIT'] = str(multiprocessing.cpu_count())

                # Define the base folder based on medium (Hindi or English)
                medium_folder = "./preuploaded"
                medium_options = ["Select Medium", "Hindi Medium", "English Medium"]
                selected_medium = st.selectbox("Select Medium", medium_options, index=0, key="medium_selector")

                if selected_medium != "Select Medium":
                    # Adjust the folder path based on the selected medium
                    language = "Hindi" if selected_medium == "Hindi Medium" else "English"
                    subjects_folder = os.path.join(medium_folder, selected_medium)
                    subjects_list = [d for d in os.listdir(subjects_folder) if os.path.isdir(os.path.join(subjects_folder, d))]
                    subjects_list.sort()
                    subjects_list.insert(0, "Select subject")

                    selected_subject = st.selectbox("Select a subject", subjects_list, index=0, key="subject_selector")

                    if selected_subject and selected_subject != "Select subject":
                        folder_path = os.path.join(subjects_folder, selected_subject)
                        files_list = list_files(folder_path)
                        files_list = [remove_extension(filename) for filename in files_list]
                        files_list.insert(0, "Select documents")

                        selected_file = st.multiselect(
                            "Select files (or select all)",
                            ["Select All"] + files_list,
                            key="terminologies_selected_files"
                        )

                        if "Select All" in selected_file:
                            selected_files = files_list[1:]  # Exclude "Select documents"
                        else:
                            selected_files = [f for f in selected_file if f != "Select documents"]

                        if selected_files:  # Check if any files are selected
                            progress_bar = st.progress(0)
                            status_text = st.empty()

                            # Initialize session state for storing results
                            if 'combined_text' not in st.session_state:
                                st.session_state.combined_text = []
                            
                            # Process PDFs and get combined text
                            combined_text, processed_files = batch_process_pdfs_with_cache(
                                selected_files,
                                folder_path,
                                progress_bar,
                                status_text
                            )

                            if combined_text:
                                status_text.text("Generating terminologies and keyterms...")
                                final_text = "\n\n".join(combined_text)

                                try:
                                    # Initialize OpenAI model with correct model name
                                    llm = ChatOpenAI(
                                        model="gpt-4o",  # Fixed model name from "gpt-4o" to "gpt-4"
                                        temperature=0.7,
                                        api_key=openai_api_key2
                                    )
                                    # Create conversation chain
                                    st.session_state.mcq_chain = ConversationChain(llm=llm)
                                    
                                    # Generate terminologies
                                    formatted_response = st.session_state.mcq_chain.predict(
                                        input=learn_outcome_prompt.format(
                                            final_text,
                                            language
                                              # Default to English if not set
                                        )
                                    )
                                    
                                    # Display the terminologies
                                    st.markdown("### Generated Terminologies and Key Terms")
                                    st.markdown(formatted_response)
                                    
                                    # Generate output files
                                    output_filename = "_".join(processed_files) if len(processed_files) <= 3 else f"{processed_files[0]}_and_{len(processed_files)-1}_others"
                                    
                                    # Create downloads
                                    markdown_to_pdf(formatted_response, f'{output_filename}_terminologies.pdf')
                                    word_doc = create_word_doc(formatted_response)
                                    doc_buffer = download_doc(word_doc)
                                    
                                    col1, col2 = st.columns(2)
                                    with col1:
                                        st.download_button(
                                            label="Download Word Document",
                                            data=doc_buffer,
                                            file_name=f"{output_filename}_terminologies.docx",
                                            mime="application/octet-stream",
                                            key='worddownload_combined'
                                        )
                                    
                                    progress_bar.empty()
                                    status_text.empty()
                                    st.success(f"Successfully processed {len(processed_files)} files and generated terminologies")
                                
                                except Exception as e:
                                    st.error(f"Error generating terminologies: {str(e)}")
                                    st.error("Please check your OpenAI API key and model access.")
                            else:
                                st.error("No text could be extracted from any of the selected files")



            if choose=="Text Analyzer":
                txt = st.text_area(
                "Text to Generate Questions"
                )
                if len(txt)>0:
                    #st.write(txt)
                    with open('dictionary.json','r') as f:
                            existing_dictionary = json.load(f)

                    lowercase_dict = {key.lower(): value for key, value in existing_dictionary.items()}
                    st.session_state.text=correct_bhashni_translations(txt,lowercase_dict)
                    #st.write(st.session_state.text)
                    with st.form(key="Text Analyzer"):
                        col1, col2 = st.columns(2)
                        with col1:
                            #st.session_state.complexity =  st.selectbox('Complexity Mode Required?*', ['No', 'Yes'],index=0,key="mode")
                            st.session_state.no_of_questions = st.number_input('No. of  Questions to generate*',key="ai_questions_no_k",step=1,max_value=30)
                            st.session_state.mode_of_questions = st.selectbox('Choose Answer Required/Not*', ['Select Option','Only Questions', 'Questions with Answers'],index=0,key="quesansw")
                        with col2:
                            st.session_state.language = st.selectbox('Choose Response Language Mode*', ['Hindi','English', 'English and Hindi'], index=0, key="lang")
                            st.session_state.type_of_questions =  st.selectbox('Choose Question Type*', ['Select Option','Short Questions','MCQ','Fill in the Blanks','True and False'],index=0,key="ai_questions_no_s")
                        submitted = st.form_submit_button("Submit")
                    if submitted and st.session_state.text and st.session_state.no_of_questions>0:
                            st.session_state.llm = ConversationChain( llm=ChatOpenAI(
                            model="gpt-4o",
                            temperature=0.7,
                            api_key=openai_api_key2
                            )) 
                            formatted_output = st.session_state.llm.predict(input = ai_prompt.format(st.session_state.no_of_questions,
                                                                            st.session_state.mode_of_questions,
                                                                            st.session_state.type_of_questions,
                                                                            st.session_state.language,
                                                                            st.session_state.text))
                            st.info(formatted_output)
                            markdown_to_pdf(formatted_output,'question.pdf')
                            word_doc = create_word_doc(formatted_output)
                            doc_buffer = download_doc(word_doc)
                            st.session_state.doc_buffer = doc_buffer
                            if 'doc_buffer' in st.session_state:
                                st.download_button(label="Download Word Document", 
                                            data=doc_buffer, 
                                            file_name="generated_document.docx", 
                                            mime="application/octet-stream",
                                            key='worddownload2')

            else:
                    st.write("")

            if choose=="Image Analyzer":
                
                #openai.api_version = ""
                #openai.api_base = ""  
                #openai.api_key = 
                #OPENAI_API_KEY2 =
                #openai_api_key2 = st.secrets["secret_section"]["OPENAI_API_KEY"]

                def load_image(img):
                    im = Image.open(img)
                    im = im.resize((400, 300))
                    image = np.array(im)
                    return image

                def to_base64(uploaded_file):
                    if uploaded_file is not None:
                       file_buffer = uploaded_file.read()
                       b64 = base64.b64encode(file_buffer).decode()
                       return f"data:image/png;base64,{b64}"
                    return None

                def generate_questions(image_base64):
                    response = openai.chat.completions.create(
                      model="gpt-4o",
                      messages=[
                          {
                            "role": "user",
                             "content": [
                                 {"type": "text", "text": st.session_state.text_prompt},
                                 {
                                       "type": "image_url",
                                       "image_url": {
                                        "url": image_base64,
                                        },
                                 },
                              ],
                           }
                    ],
                    max_tokens=2000,
                )
                    return response.choices[0].message.content

                st.write("Upload Image to Generate Questions")
                st.session_state.image = st.file_uploader(label=" ", accept_multiple_files=False, type=["jpg", "jpeg", "png"])

                if st.session_state.image:
                   col1, col2 = st.columns(2)
                   with col1:
                        st.session_state.complexity = st.selectbox('Complexity Mode Required?*', ['Easy', 'Difficult'], index=0, key="mode1image")
                        st.session_state.no_of_questions = st.number_input('No. of Questions to generate*', key="ai_questions_no_a_image", step=1, max_value=30)
                        st.session_state.mode_of_questions = st.selectbox('Choose Answer Required/Not*', ['Select Option', 'Only Questions', 'Questions with Answers'], index=0, key="quesanswz_image")
                   with col2:
                        st.session_state.type_of_questions = st.selectbox('Choose Question Type*', ['Select Option', 'Short Questions', 'MCQ', 'Fill in the Blanks', 'True and False'], index=0, key="ai_questions_no_p_image")
                        st.session_state.classq = st.selectbox('Choose Class*', ['Select Option', '1', '2', '3', '4', '5', '6', '7', '8', '9', '10', '11', '12'], index=0, key="ai_questions_no_p1_image")
                        st.session_state.cont = st.text_input("Any other context (Optional)", key="eng_image")

                   if st.session_state.mode_of_questions != "Select Option":
                       st.session_state.text_prompt = f'''Based on the image, generate only questions considering following constraints,
                       1. number of questions  - {st.session_state.no_of_questions}
                       2. mode of questions - {st.session_state.mode_of_questions}
                       3. type of questions - {st.session_state.type_of_questions}
                       4. Level of questions - {st.session_state.complexity}
                       5. Class - {st.session_state.classq}
                       6. Question Context - {st.session_state.cont}
                       Generate questions according to Madhya Pradesh School Education Board
                       Response is to be generated in both English and Hindi, first generate in English then in Hindi
                       after generate Answer should be start new line.
                       '''
        
                       image_base64 = to_base64(st.session_state.image)
                       if image_base64:
                          st.write("Image successfully converted to base64")
                          formatted_output = generate_questions(image_base64)
                          img = load_image(st.session_state.image)
                          st.image(img)
                          st.info(formatted_output)
                       else:
                          st.error("Failed to convert image to base64.")
                else:
                    st.info("Please upload an image file.")
   
                
            if choose == "Topic Based Questions":
                with st.form(key="Topic Based Questions"):
                    col1, col2 = st.columns(2)
                    with col1:
                        st.session_state.topic_name = st.text_input('Specific Topic Name', placeholder="Topic Name", key="tt")
                        st.session_state.complexity = st.selectbox('Complexity Mode Required?*', ['Easy', 'Difficult'], index=0, key="mode1")
                        st.session_state.no_of_questions = st.number_input('No. of Questions to generate*', key="ai_questions_no_a", step=1, max_value=30)
                    with col2:
                        st.session_state.type_of_questions = st.selectbox('Choose Question Type*', ['Select Option', 'Short Questions', 'MCQ', 'Fill in the Blanks', 'True and False'], index=0, key="ai_questions_no_p")
                        st.session_state.language = st.selectbox('Choose Response Language Mode*', ['Hindi','English', 'English and Hindi'], index=0, key="lang")
                        st.session_state.mode_of_questions = st.selectbox('Choose Answer Required/Not*', ['Select Option', 'Only Questions', 'Questions with Answers'], index=0, key="quesanswz")

                    submitted = st.form_submit_button("Submit")

                if submitted and st.session_state.topic_name and st.session_state.mode_of_questions != 'Select Option':
                    st.session_state.llm = ConversationChain(llm=ChatOpenAI(
                        model="gpt-4o-mini",
                        temperature=0.7,
                        api_key=openai_api_key2
                    ))

                    formatted_output = st.session_state.llm.predict(input=ai_topic_prompt.format(
                        st.session_state.topic_name,
                        st.session_state.no_of_questions,
                        st.session_state.mode_of_questions,
                        st.session_state.type_of_questions,
                        st.session_state.language,
                        st.session_state.complexity
                    ))

                    st.write(formatted_output)

                    # Generate PDF and Word files
                    markdown_to_pdf(formatted_output, 'question.pdf')
                    word_doc = create_word_doc(formatted_output)
                    doc_buffer = download_doc(word_doc)

                    # Move the download button outside the form
                    st.download_button(label="Download Word Document",
                                    data=doc_buffer,
                                    file_name="generated_document.docx",
                                    mime="application/octet-stream",
                                    key='worddownload3')
            if choose == "Examination":
                with st.form(key="Examination"):
                    col1, col2 = st.columns(2)
                    with col1:
                        unit_name = st.text_input('Unit Name', placeholder="Enter Unit Name", key="unit_name")
                        topic_name = st.text_input('Specific Topic Name', placeholder="Topic Name", key="tt")
                        complexity = st.selectbox('Complexity Mode Required?*', ['Easy', 'Difficult'], index=0, key="mode1")
                        no_of_questions = st.number_input('No. of Questions to generate*', key="ai_questions_no_a", step=1, max_value=30)
                    with col2:
                        type_of_questions = st.selectbox('Choose Question Type*', ['Select Option', 'Short Questions', 'MCQ', 'Fill in the Blanks', 'True and False'], index=0, key="ai_questions_no_p")
                        language = st.selectbox('Choose Response Language Mode*', ['Hindi', 'English', 'English and Hindi'], index=0, key="lang")
                        mode_of_questions = st.selectbox('Choose Answer Required/Not*', ['Select Option', 'Only Questions', 'Questions with Answers'], index=0, key="quesanswz")
                        level = st.selectbox('Select Question Level', ['All','L1', 'L2', 'L3', 'L4'], index=0, key="level")

                    submitted = st.form_submit_button("Submit")

                if submitted and unit_name and topic_name and mode_of_questions != 'Select Option':
                    # Show a loading spinner while generating questions
                    with st.spinner('Generating questions...'):
                        llm = ConversationChain(llm=ChatOpenAI(
                            model="gpt-4o-mini",
                            temperature=0.7,
                            api_key=openai_api_key2
                        ))

                        # Modified prompt for clean output
                        ai_prompt = """Generate {0} {1} {2} questions about {3} with complexity level {4} in {5} language. 

            Follow this exact format for each question:

            [Question text]
            a) [Option 1]
            b) [Option 2]
            c) [Option 3]
            d) [Option 4]
            Answer: [correct option letter]

            Example:
            What is the capital of France?
            a) London
            b) Paris
            c) Berlin
            d) Madrid
            Answer: b

            Important formatting rules:
            - Write questions directly without any numbers or prefixes
            - Start options with a), b), c), d)
            - Include Answer: followed by the correct option letter
            - Separate questions with one blank line
            - No additional text or explanations

            Generate the questions now:"""

                        formatted_output = llm.predict(input=ai_prompt.format(
                            no_of_questions,
                            complexity,
                            type_of_questions,
                            topic_name,
                            level,
                            language
                        ))

                        st.write(formatted_output)

                        # Process the output into structured data
                        question_data = []
                        
                        # Split into individual questions
                        questions = [q.strip() for q in formatted_output.split('\n\n') if q.strip()]
                        
                        for question in questions:
                            try:
                                # Parse the question
                                question_parts = parse_question(question)
                                
                                # Only add if we have all required parts and parsing was successful
                                if question_parts and all(value is not None for value in question_parts.values()):
                                    question_data.append({
                                        "Unit": unit_name,
                                        "Level": level,
                                        "Topic": topic_name,
                                        **question_parts
                                    })
                            except Exception as e:
                                st.error(f"Error processing question: {str(e)}")
                                continue

                        # Generate Excel file if we have questions
                        if question_data:
                            excel_buffer = create_excel(question_data)
                            
                            st.download_button(
                                label="Download Excel Document",
                                data=excel_buffer,
                                file_name="Ict_Question_Bank.xlsx",
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                key='excel_download'
                            )

                        else:
                            st.error("No valid questions were generated. Please try again.")

            if choose == "Question Paper Generator":
                os.environ['OMP_THREAD_LIMIT'] = str(multiprocessing.cpu_count())

                # Define the base folder based on medium
                medium_folder = "./preuploaded"
                medium_options = ["Select Medium", "Hindi Medium", "English Medium"]
                selected_medium = st.selectbox("Select Medium", medium_options, index=0, key="medium_selector")

                if selected_medium != "Select Medium":
                    # Adjust the folder path based on the selected medium
                    subjects_folder = os.path.join(medium_folder, selected_medium)
                    subjects_list = [d for d in os.listdir(subjects_folder) if os.path.isdir(os.path.join(subjects_folder, d))]
                    subjects_list.sort()
                    subjects_list.insert(0, "Select subject")

                    selected_subject = st.selectbox("Select a subject", subjects_list, index=0, key="subject_selector")

                    if selected_subject and selected_subject != "Select subject":
                        folder_path = os.path.join(subjects_folder, selected_subject)
                        files_list = list_files(folder_path)
                        files_list = [remove_extension(filename) for filename in files_list]
                        files_list.insert(0, "Select documents")

                        selected_file = st.multiselect(
                            "Select files (or select all)",
                            ["Select All"] + files_list,
                            key="qp_selected_files"
                        )

                        if "Select All" in selected_file:
                            selected_files = files_list[1:]  # Exclude "Select documents"
                        else:
                            selected_files = [f for f in selected_file if f != "Select documents"]

                        if selected_files:
                            progress_bar = st.progress(0)
                            status_text = st.empty()

                            # Process PDFs and get combined text
                            combined_text, processed_files = batch_process_pdfs_with_cache(
                                selected_files,
                                folder_path,
                                progress_bar,
                                status_text
                            )

                            if combined_text:
                                status_text.text("Processing content for question paper generation...")
                                final_text = "\n\n".join(combined_text)

                                # Question Paper Generation Form
                                with st.form(key="question_paper_form"):
                                    
                                    
                                    terminal = st.selectbox(
                                        'Exam Terminal*', 
                                        ['1st Term', '2nd Term', 'Pre Board', 'Unit Test'],
                                        key="terminal"
                                    )
                                
                                
                                    # Set language based on selected medium
                                    language = "Hindi" if selected_medium == "Hindi Medium" else "English"
                                    #st.text(f"Selected Language: {language}")

                                    submit_button = st.form_submit_button("Generate Question Paper")

                                if submit_button:
                                    try:
                                        status_text.text("Generating question paper...")
                                        
                                        # Initialize OpenAI model with correct model name
                                        llm = ChatOpenAI(
                                            model="gpt-4o",  # Fixed model name
                                            temperature=0.7,
                                            api_key=openai_api_key2
                                        )
                                        
                                        chain = ConversationChain(llm=llm)
                                        
                                        # Format chapter information
                                        chapters = ", ".join(selected_files)
                                        
                                        # Generate question paper using the updated prompt
                                        formatted_output = chain.predict(
                                            input=ai_topic_prompt2.format(
                                                final_text,  # Terminologies and content
                                                language,    # Selected language
                                                chapters     # Chapter information
                                            )
                                        )

                                        # Display generated question paper
                                        st.markdown("### Generated Question Paper")
                                        st.markdown(formatted_output)

                                        # Generate output files
                                        output_filename = f"{selected_subject}_{terminal}_question_paper"
                                        
                                        # Create downloads
                                        markdown_to_pdf(formatted_output, f'{output_filename}.pdf')
                                        word_doc = create_word_doc(formatted_output)
                                        doc_buffer = download_doc(word_doc)
                                        
                                        # Download buttons
                                        col1, col2 = st.columns(2)
                                        with col1:
                                            st.download_button(
                                                label="Download Word Document",
                                                data=doc_buffer,
                                                file_name=f"{output_filename}.docx",
                                                mime="application/octet-stream",
                                                key='worddownload'
                                            )
                                        
                                        progress_bar.empty()
                                        status_text.empty()
                                        st.success("Question paper generated successfully!")
                                        
                                    except Exception as e:
                                        st.error(f"Error generating question paper: {str(e)}")
                                        st.error("Please check your OpenAI API key and model access.")
                                
                            else:
                                st.error("No text could be extracted from the selected files")

                                        
if st.session_state.teach=='Students':
    choose=st.radio("Select Options",("Pre Uploaded","Ask a Query","Text Analyzer","Student Exam Attempt"),horizontal=True)
    if choose == "Pre Uploaded":
        
        os.environ['OMP_THREAD_LIMIT'] = str(multiprocessing.cpu_count())

        # Define the base folder based on medium (Hindi or English)
        medium_folder = "./preuploaded"
        medium_options = ["Select Medium", "Hindi Medium", "English Medium"]
        selected_medium = st.selectbox("Select Medium", medium_options, index=0, key="medium_selector")

        if selected_medium != "Select Medium":
            # Adjust the folder path based on the selected medium
            subjects_folder = os.path.join(medium_folder, selected_medium)
            subjects_list = [d for d in os.listdir(subjects_folder) if os.path.isdir(os.path.join(subjects_folder, d))]
            subjects_list.sort()
            subjects_list.insert(0, "Select subject")

            selected_subject = st.selectbox("Select a subject", subjects_list, index=0, key="subject_selector")

            if selected_subject and selected_subject != "Select subject":
                folder_path = os.path.join(subjects_folder, selected_subject)
                files_list = list_files(folder_path)
                files_list = [remove_extension(filename) for filename in files_list]
                files_list.insert(0, "Select documents")

                selected_file = st.multiselect(
                    "Select files (or select all)",
                    ["Select All"] + files_list,
                    key="terminologies_selected_files"
                )

                if "Select All" in selected_file:
                    selected_files = files_list[1:]
                elif "Select documents" in selected_file:
                    selected_file.remove("Select documents")

                vector_store = initialize_chroma()  # Replace with your initialization logic

                # Initialize combined_text if it doesn't exist
                if "combined_text" not in st.session_state:
                    st.session_state.combined_text = None
                
                if selected_file:
                    progress_bar = st.progress(0)
                    status_text = st.empty()

                    combined_text, processed_files = batch_process_pdfs_with_cache(
                        selected_file,
                        folder_path,
                        progress_bar,
                        status_text
                    )
                    if combined_text:
                        status_text.text("Generating terminologies and keyterms...")
                        st.session_state.final_text = "\n\n".join(combined_text)

                    # Handling form for question generation
                    with st.form(key="Pre Uploaded"):
                        col1, col2 = st.columns(2)
                        with col1:
                            st.session_state.complexity = st.selectbox('Complexity Mode Required?*', ['Easy', 'Difficult'], index=0, key="mode")
                            st.session_state.no_of_questions = st.number_input('No. of Questions to generate*', key="ai_questions", step=1, max_value=30)
                            st.session_state.mode_of_questions = st.selectbox('Choose Answer Required/Not*', ['Only Questions', 'Questions with Answers'], index=0, key="quesansw")
                        with col2:
                            st.session_state.type_of_questions = st.selectbox('Choose Question Type*', ['Short Questions', 'Long Questions', 'MCQ', 'Fill in the Blanks', 'True and False'], index=0)
                        submitted = st.form_submit_button("Submit")

                    if submitted and st.session_state.final_text and st.session_state.mode_of_questions != 'Select Option':
                        if st.session_state.final_text:
                            st.session_state.llm = ConversationChain(llm=ChatOpenAI(model="gpt-4o", temperature=0.7, api_key=openai_api_key2))
                            chapter_info = f"Chapter: {selected_file}" if selected_file != "All Chapters" else "All Chapters"

                            # Determine language based on medium
                            language = "Hindi" if selected_medium == "Hindi" else "English"

                            # Generate the formatted output based on selected options
                            formatted_output = st.session_state.llm.predict(input=ai_topic_prompt1.format(
                                chapter_info,
                                st.session_state.no_of_questions,
                                st.session_state.final_text,
                                language,
                                st.session_state.mode_of_questions,
                                st.session_state.type_of_questions,
                                st.session_state.complexity,
                                st.session_state.no_of_questions
                            ))

                            st.info(formatted_output)
                            markdown_to_pdf(formatted_output, 'question.pdf')
                            word_doc = create_word_doc(formatted_output)
                            doc_buffer = download_doc(word_doc)

                            st.download_button(
                                label="Download Word Document",
                                data=doc_buffer,
                                file_name="generated_document.docx",
                                mime="application/octet-stream",
                                key='worddownload'
                            )
                        else:
                            st.info("No relevant results found based on the subject and chapter metadata.") 

# Main app logic


# Helper function to read the Excel file and extract questions based on type
    if choose == "Student Exam Attempt":
        subjects_folder = "./preuploaded"

        # List subjects
        subjects_list = [d for d in os.listdir(subjects_folder) if os.path.isdir(os.path.join(subjects_folder, d))]
        subjects_list.sort()
        subjects_list.insert(0, "Select subject")

        selected_subject = st.selectbox("Select a subject", subjects_list, index=0, key='subject_selector')

        if selected_subject != "Select subject":
            # List chapters/documents for the selected subject
            folder_path = os.path.join(subjects_folder, selected_subject)
            files_list = list_files(folder_path)
            files_list = [remove_extension(filename) for filename in files_list]
            files_list.insert(0, "Select document")

            selected_file = st.selectbox("Select a file", files_list, index=0, key='chapter_selector')

            # Language and exam terminal selection
            col1, col2 = st.columns(2)
            with col1:
                st.session_state.Terminal = st.selectbox('Exam Terminal*', ['1st Term', '2nd Term', 'Pre Board', 'Unit Test'], index=0, key="mode")
            with col2:
                st.session_state.language = st.selectbox('Choose Response Language Mode*', ['Hindi', 'English', 'English and Hindi'], index=0, key="lang")

            if selected_file != "Select document":
                # Initialize Chroma DB
                vector_store = initialize_chroma()

                # Initialize filename for the selected document
                st.session_state.filename = []
                pdf_file_path = os.path.join(folder_path, selected_file + '.pdf')

                # Extract text from the selected PDF file
                text = pdf_to_text(pdf_file_path)

                st.session_state.filename.append(selected_file)

                # Check if embeddings already exist for the document
                if not check_embeddings_exist(vector_store, selected_subject, selected_file):
                # Create and store embeddings only if they don't exist
                    create_and_store_embeddings(text, selected_subject, selected_file, vector_store)
                else:
                    st.info("Embeddings for this document already exist in the database.")

                # Helper function to generate the question paper
                def generate_question_paper(text_content: str) -> str:
                    # Pass the text and language selection to the LLM to generate the question paper
                    formatted_response = st.session_state.llm.predict(
                        input=ai_topic_prompt3.format(f"Chapter: {selected_file}" if selected_file != "All Chapters" else "All Chapters", st.session_state.language)
                    )
                    return formatted_response

                # Initialize LLM for generating the question paper
                st.session_state.llm = ConversationChain(
                    llm=ChatOpenAI(model="gpt-4o", temperature=0.7, api_key=openai_api_key2)
                )

                # Function to extract questions from the generated output
                def extract_questions(generated_output):
                    questions = []
                    current_question = ""
                    for line in generated_output.split('\n'):
                        line = line.strip()
                        if line.startswith('Q') and line[1].isdigit():
                            if current_question:
                                questions.append(current_question.strip())
                            current_question = line
                        elif current_question and not line.startswith(('Objective Section', 'Subjective Section', 'Multiple Choice Questions', 'Fill in the Blanks', 'True/False Statements', 'One-word Answers')):
                            current_question += '\n' + line
                    if current_question:
                        questions.append(current_question.strip())
                    return questions

                # Use st.form to prevent re-running after every user interaction
                # Use st.form to prevent re-running after every user interaction
                with st.form("exam_form"):
                    st.subheader("Please attempt the following questions:")

                    # Extract questions from the generated output
                    generated_output = generate_question_paper(text)
                    questions = extract_questions(generated_output)

                    # Initialize a dictionary to hold user answers
                    user_answers = {}

                    # Loop through each question and create an input field below it
                    for index, question in enumerate(questions):
                        st.markdown(f"**Q{index + 1}:** {question}")

                        # Create a text area for each question with clear instruction
                        user_answer = st.text_area(f"Your answer for Q{index + 1} (write your answer below):", height=100, key=f"answer_{index}")
                        user_answers[f"Q{index + 1}"] = user_answer  # Store the answer in the dictionary

                    # Submit button for the form
                    submitted = st.form_submit_button("Submit Answers")

                if submitted:
                    st.success("Your answers have been submitted!")

                    # Create a list to store question data
                    question_data = []

                    # Process user answers and add to question_data
                    for index, (question, answer) in enumerate(zip(questions, user_answers.values())):
                    # You may want to determine this dynamically
                        question_data.append({
                            
                            "Question": question,
                            "Response": answer
                        })

                    # Create a DataFrame from the question data
                    df = pd.DataFrame(question_data)

                    # Define and create terminal folder **before** using it
                    base_folder = './exam_uploaded'
                    terminal_folder = os.path.join(base_folder, st.session_state.Terminal)
                    if not os.path.exists(base_folder):
                        os.makedirs(base_folder)
                    if not os.path.exists(terminal_folder):
                        os.makedirs(terminal_folder)

                    # Save the DataFrame to an Excel file
                    excel_file_path = os.path.join(terminal_folder, f"{selected_subject}_{selected_file}_responses.xlsx")
                    df.to_excel(excel_file_path, index=False)

                    st.success(f"Responses saved to Excel file: {excel_file_path}")

                    # Save the generated output as a PDF and Word document
                    markdown_to_pdf(generated_output, 'question_paper.pdf')
                    word_doc = create_word_doc(generated_output)
                    doc_buffer = download_doc(word_doc)

                    # Provide a download button for the Word document
                    st.download_button(
                        label="Download Word Document",
                        data=doc_buffer,
                        file_name="question_paper.docx",
                        mime="application/octet-stream",
                        key='worddownload'
                    )

                    # Save PDF and Word files
                    pdf_file_path = os.path.join(terminal_folder, 'question_paper.pdf')
                    word_file_path = os.path.join(terminal_folder, 'generated_document.docx')
                    word_doc.save(word_file_path)

                    # Upload success feedback
                    #st.success(f"{st.session_state.Terminal} Question paper submitted successfully!")
                    #st.success(f"Responses saved to Excel file: {excel_file_path}")



    if choose=="Text Analyzer":
                
                txt = st.text_area(
                "Text to Generate Questions"
                )
                if len(txt)>0:
                    #st.write(txt)
                    with open('dictionary.json','r') as f:
                            existing_dictionary = json.load(f)

                    lowercase_dict = {key.lower(): value for key, value in existing_dictionary.items()}
                    st.session_state.text=correct_bhashni_translations(txt,lowercase_dict)
                    #st.write(st.session_state.text)

                    with st.form(key="Text Analyzer"):
                        col1, col2 = st.columns(2)
                        with col1:
                            #st.session_state.complexity =  st.selectbox('Complexity Mode Required?*', ['No', 'Yes'],index=0,key="mode")
                            st.session_state.no_of_questions = st.number_input('No. of  Questions to generate*',key="ai_questions_no_k",step=1,max_value=30)
                            st.session_state.mode_of_questions = st.selectbox('Choose Answer Required/Not*', ['Select Option','Only Questions', 'Questions with Answers'],index=0,key="quesansw")
                        with col2:
                            st.session_state.language = st.selectbox('Choose Response Language Mode*', ['Hindi','English', 'English and Hindi'], index=0, key="lang")
                            st.session_state.type_of_questions =  st.selectbox('Choose Question Type*', ['Select Option','Short Questions','MCQ','Fill in the Blanks','True and False'],index=0,key="ai_questions_no_s")
                        submitted = st.form_submit_button("Submit")
                    if submitted and st.session_state.text and st.session_state.no_of_questions>0:
                            st.session_state.llm = ConversationChain( llm=ChatOpenAI(
                            model="gpt-4o",
                            temperature=0.7,
                            api_key=openai_api_key2
                            )) 
                            formatted_output = st.session_state.llm.predict(input = ai_prompt.format(st.session_state.no_of_questions,
                                                                            st.session_state.mode_of_questions,
                                                                            st.session_state.type_of_questions,
                                                                            st.session_state.language,
                                                                            st.session_state.text))
                            st.info(formatted_output)
                            markdown_to_pdf(formatted_output,'question.pdf')
                            word_doc = create_word_doc(formatted_output)
                            doc_buffer = download_doc(word_doc)
                            st.session_state.doc_buffer = doc_buffer
                            if 'doc_buffer' in st.session_state:
                                st.download_button(label="Download Word Document", 
                                            data=doc_buffer, 
                                            file_name="generated_document.docx", 
                                            mime="application/octet-stream",
                                            key='worddownload2')

    else:
        st.write("")



    # Inside call_chatgpt function
    def call_chatgpt(prompt, model_name):
        return prompts.SYSTEM_PROMPTS["Default"]
    
    if choose=="MCQ Test":
        st.write("In Development Stage")
        st.write('Note: File name should contain subject and class like maths_class10.pdf/.docx')
        files = st.file_uploader('Upload Books,Notes,Question Banks ', accept_multiple_files=True,type=['pdf', 'docx'])
        if files:
            file_extension = files[0].name.split(".")[-1]
            if file_extension == "pdf":
                path = files[0].read()
                name=files[0].name[:-4]
                # Check if the file exists
                if not os.path.exists(name+".txt"):
                    print("File Not Exist")
                    with open(name+'.txt', 'w') as file:
                    # Open a file in write mode (creates a new file if it doesn't exist)
                        st.session_state.text= chat.load_pdf_text(files[0],name)
                        file.write(st.session_state.text)
                        # st.session_state.mcq_chain = ConversationChain( llm=AzureChatOpenAI(
                        # deployment_name="gpt-4turbo",
                        # temperature=0
                        # ))
                        #outputs = chat.mcq_response(st.session_state.text)
                        outputs = generate_quiz(st.session_state.text)
                        #st.info(formatted_output)
                        try:
                            quiz_json = json.loads(outputs[0])["quiz"]
                            st.session_state['quiz_length'] = len(quiz_json)
                            questions = [q['question'] for q in quiz_json]
                            options = [q['options'] for q in quiz_json]
                            answers = [q["answer"] for q in quiz_json]
                            explanations = [q['explanation'] for q in quiz_json]
                            
                            # user selects which question they want to answer
                            # question_num = st.number_input('Choose a question', min_value=1, max_value = len(quiz_json), value=1)
                            question_num = st.session_state['curr_question']
                            answer_choices = options[question_num]
                            col1, col2 = st.columns(2)
                            with col1:
                                st.button('Previous Question', use_container_width=True, on_click=decrement_question_num)
                            with col2:
                                st.button('Next Question', use_container_width=True, on_click=increment_question_num)

                            st.markdown(f"##### Question {question_num + 1} of {len(quiz_json)}: {questions[question_num]}")

                            if 'a' not in st.session_state:
                                st.session_state.a = 0
                                st.session_state.b = 0
                                st.session_state.c = 0
                                st.session_state.d = 0

                            def ChangeA():
                                st.session_state.a,st.session_state.b,st.session_state.c,st.session_state.d = 1,0,0,0
                            def ChangeB():
                                st.session_state.a,st.session_state.b,st.session_state.c,st.session_state.d = 0,1,0,0
                            def ChangeC():
                                st.session_state.a,st.session_state.b,st.session_state.c,st.session_state.d = 0,0,1,0
                            def ChangeD():
                                st.session_state.a,st.session_state.b,st.session_state.c,st.session_state.d = 0,0,0,1
                            def ClearAll():
                                st.session_state.a,st.session_state.b,st.session_state.c,st.session_state.d = 0,0,0,0

                            checkboxA = st.checkbox(answer_choices[0], value = st.session_state.a, on_change = ChangeA)
                            checkboxB = st.checkbox(answer_choices[1], value = st.session_state.b, on_change = ChangeB)
                            checkboxC = st.checkbox(answer_choices[2], value = st.session_state.c, on_change = ChangeC)
                            checkboxD = st.checkbox(answer_choices[3], value = st.session_state.d, on_change = ChangeD)

                            if st.session_state.a:
                                user_answer = answer_choices[0]
                            elif st.session_state.b:
                                user_answer = answer_choices[1]
                            elif st.session_state.c:
                                user_answer = answer_choices[2]
                            elif st.session_state.d:
                                user_answer = answer_choices[3]
                            else:
                                user_answer = None

                            if user_answer is not None:
                                user_answer_num = answer_choices.index(user_answer)
                                if st.button('Submit Answer', type='secondary'):
                                    if user_answer_num == answers[question_num][0]:
                                        st.success(f'Correct! {explanations[question_num]}')
                                    else:
                                        st.error(f'Incorrect :( \n\n The correct answer was: {answer_choices[answers[question_num][0]]}\n\n {explanations[question_num]}')
                        except:
                            st.info('Uh oh... could not generate a quiz for ya! Happy studying!')
                else:
                    print("file Exist")
                    st.session_state.filename=[]
                    with open(name+'.txt', 'r',encoding='ISO-8859-1') as file:
            #        Read the content of the file
                        st.session_state.filename.append(name)
                        st.session_state.text = file.read()
                        #st.write(st.session_state.text)
                        #outputs = generate_quiz(st.session_state.text)
                        # st.session_state.mcq_chain = ConversationChain( llm=AzureChatOpenAI(
                        # deployment_name="gpt-4turbo",
                        # temperature=0
                        # ))
                        #outputs = chat.mcq_response(st.session_state.text)
                        outputs = generate_quiz(st.session_state.text)
                        #st.write(outputs)
                        #print("Output are")
                        #print(outputs)
                        
                        st.write(outputs)
                        corrected_json_data = outputs[0].replace('//', '')
                        quiz_json = json.loads(corrected_json_data)["quiz"]
                        st.session_state['quiz_length'] = len(quiz_json)
                        questions = [q['question'] for q in quiz_json]
                        options = [q['options'] for q in quiz_json]
                        answers = [q["answer"] for q in quiz_json]
                        explanations = [q['explanation'] for q in quiz_json]
                        
                        # user selects which question they want to answer
                        # question_num = st.number_input('Choose a question', min_value=1, max_value = len(quiz_json), value=1)
                        question_num = st.session_state['curr_question']
                        answer_choices = options[question_num]
                        col1, col2 = st.columns(2)
                        with col1:
                            st.button('Previous Question', use_container_width=True, on_click=decrement_question_num)
                        with col2:
                            st.button('Next Question', use_container_width=True, on_click=increment_question_num)

                        st.markdown(f"##### Question {question_num + 1} of {len(quiz_json)}: {questions[question_num]}")

                        if 'a' not in st.session_state:
                            st.session_state.a = 0
                            st.session_state.b = 0
                            st.session_state.c = 0
                            st.session_state.d = 0

                        def ChangeA():
                            st.session_state.a,st.session_state.b,st.session_state.c,st.session_state.d = 1,0,0,0
                        def ChangeB():
                            st.session_state.a,st.session_state.b,st.session_state.c,st.session_state.d = 0,1,0,0
                        def ChangeC():
                            st.session_state.a,st.session_state.b,st.session_state.c,st.session_state.d = 0,0,1,0
                        def ChangeD():
                            st.session_state.a,st.session_state.b,st.session_state.c,st.session_state.d = 0,0,0,1
                        def ClearAll():
                            st.session_state.a,st.session_state.b,st.session_state.c,st.session_state.d = 0,0,0,0

                        checkboxA = st.checkbox(answer_choices[0], value = st.session_state.a, on_change = ChangeA)
                        checkboxB = st.checkbox(answer_choices[1], value = st.session_state.b, on_change = ChangeB)
                        checkboxC = st.checkbox(answer_choices[2], value = st.session_state.c, on_change = ChangeC)
                        checkboxD = st.checkbox(answer_choices[3], value = st.session_state.d, on_change = ChangeD)

                        if st.session_state.a:
                            user_answer = answer_choices[0]
                        elif st.session_state.b:
                            user_answer = answer_choices[1]
                        elif st.session_state.c:
                            user_answer = answer_choices[2]
                        elif st.session_state.d:
                            user_answer = answer_choices[3]
                        else:
                            user_answer = None

                        if user_answer is not None:
                            user_answer_num = answer_choices.index(user_answer)
                            if st.button('Submit Answer', type='secondary'):
                                if user_answer_num == answers[question_num][0]:
                                    st.success(f'Correct! {explanations[question_num]}')
                                else:
                                    st.error(f'Incorrect :( \n\n The correct answer was: {answer_choices[answers[question_num][0]]}\n\n {explanations[question_num]}')


# Function to split text into smaller chunks


    openai_api_key = os.getenv("OPENAI_API_KEY")
    openai_api_key2 = st.secrets["secret_section"]["OPENAI_API_KEY"]
    if 'history' not in st.session_state:
        st.session_state.history = []

    col_1, col_2 = st.columns([2, 1])

    with col_1:
        if choose == "Ask a Query":
            uploaded_file = st.file_uploader("Upload a PDF file", type=["pdf"])

            if uploaded_file is not None:
                pdfreader = PdfReader(uploaded_file)
                raw_text = ''
                for i, page in enumerate(pdfreader.pages):
                    content = page.extract_text()
                    if content:
                        raw_text += content

                text_splitter = CharacterTextSplitter(
                    separator="\n",
                    chunk_size=800,
                    chunk_overlap=200,
                    length_function=len,
                )
                texts = text_splitter.split_text(raw_text)
                #st.write(f"PDF loaded and split into {len(texts)} chunks.")

                embeddings = OpenAIEmbeddings(api_key=openai_api_key2)
                document_search = FAISS.from_texts(texts, embeddings)
                #st.write("Document embeddings created and stored in FAISS index.")

                chain = load_qa_chain(OpenAI(api_key=openai_api_key2), chain_type="stuff")

                query = st.chat_input("Ask a question about the PDF:")
                st.write(query)
                if query:
                    docs = document_search.similarity_search(query)
                    answer = chain.run(input_documents=docs, question=query)
                    st.session_state.history.append((query, answer))
                    st.write("Answer:", answer)

                    # Translate to Hindi
                    translator = Translator(to_lang="hi")
                    
                    # Split query and answer into smaller chunks for translation
                    query_chunks = split_text_into_chunks(query, 500)
                    query_hindi_chunks = [translator.translate(chunk) for chunk in query_chunks]
                    query_hindi = " ".join(query_hindi_chunks)

                    answer_chunks = split_text_into_chunks(answer, 500)
                    answer_hindi_chunks = [translator.translate(chunk) for chunk in answer_chunks]
                    answer_hindi = " ".join(answer_hindi_chunks)

                    st.session_state.history[-1] += (query_hindi, answer_hindi)
                    st.write("**In Hindi:**")
                    st.write(f"**Q:** {query_hindi}")
                    st.write(f"**A:** {answer_hindi}")

                if st.session_state.history:
                    doc = Document()
                    doc.add_heading('Questions and Answers', 0)

                    for i, (question, answer, question_hindi, answer_hindi) in enumerate(st.session_state.history):
                        doc.add_heading(f"Q{i+1}: {question}", level=1)
                        doc.add_paragraph(f"A{i+1}: {answer}")
                        doc.add_heading(f"Q{i+1} (Hindi): {question_hindi}", level=1)
                        doc.add_paragraph(f"A{i+1} (Hindi): {answer_hindi}")

                    # Save the document in the current directory with a unique name
                    timestamp = time.strftime("%Y%m%d-%H%M%S")
                    doc_path = f"QnA_History_{timestamp}.docx"
                    doc.save(doc_path)

                    with open(doc_path, "rb") as f:
                        st.download_button(
                            label="Download Word Document",
                            data=f,
                            file_name=doc_path,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

            # Add a section for predefined prompts
            

            with col_2:
                st.write("### History")
                for i, (question, answer, question_hindi, answer_hindi) in enumerate(st.session_state.history):
                    st.write(f"**Q{i+1}:** {question}")
                    st.write(f"**A{i+1}:** {answer}")
                    st.write(f"**Q{i+1} (Hindi):** {question_hindi}")
                    st.write(f"**A{i+1} (Hindi):** {answer_hindi}")

# Add a download button to download the history as a Word document
    
    if choose=="Terminologies and Keyterms":
        st.write('Note: File name should contain subject and class like maths_class10.pdf/.docx')
        files = st.file_uploader('Upload Books,Notes,Question Banks ', accept_multiple_files=True,type=['pdf', 'docx'])
        if files:
            file_extension = files[0].name.split(".")[-1]
            if file_extension == "pdf":
                path = files[0].read()
                name=files[0].name[:-4]
                #Check if the file exists
                if not os.path.exists(name+".txt"):
                    print("File Not Exist")
                    with open(name+'.txt', 'w') as file:
                    #Open a file in write mode (creates a new file if it doesn't exist)
                        st.session_state.text= chat.load_pdf_text(files[0],name)
                        file.write(st.session_state.text)
                else:
                    print("file Exist")
                    st.session_state.filename=[]
                    with open(name+'.txt', 'r',encoding='ISO-8859-1') as file:
                    #Read the content of the file
                        st.session_state.filename.append(name)
                        st.session_state.text = file.read()
                        st.session_state.mcq_chain = ConversationChain( llm=ChatOpenAI(
                        model="gpt-3.5-turbo",
                        temperature=0.7
                        ))
                        outputs = chat.mcq_response(st.session_state.text)
                        st.write(outputs)
                        markdown_to_pdf(outputs,'question.pdf')
                        
                        word_doc = create_word_doc(outputs)
                        doc_buffer = download_doc(word_doc)
                        st.download_button(label="Download Word Document", 
                                        data=doc_buffer, 
                                        file_name="generated_document.docx", 
                                        mime="application/octet-stream",
                                        key='worddownload3')
if st.session_state.teach == 'Administration':
    if 'selected_file' not in st.session_state:
            st.session_state.selected_file = None

    # Define the options and handle the session state for the selected option
    if 'selected_option' not in st.session_state:
        st.session_state.selected_option = None

    def reset_selected_file():
        st.session_state.selected_file = None

    st.session_state.selected_option = st.radio("Select Options", ("Add Document", "Download Document", "Delete Document", "View Documents"), horizontal=True, on_change=reset_selected_file)

    choose = st.session_state.selected_option

    base_folder = "./preuploaded"

    if choose == "Add Document":
        subject_folders = [d for d in os.listdir(base_folder) if os.path.isdir(os.path.join(base_folder, d))]
        subject_folders.sort()
        subject_folders.insert(0, "Select Subject")
        
        selected_subject = st.selectbox("Select Subject", subject_folders, index=0, key='subject_selector')

        if selected_subject != "Select Subject":
            files = st.file_uploader('Upload Books, Notes, Question Banks', accept_multiple_files=True, type=['pdf'])
            if files:
                subject_folder_path = os.path.join(base_folder, selected_subject)
                os.makedirs(subject_folder_path, exist_ok=True)
                
                for uploaded_file in files:
                    file_path = os.path.join(subject_folder_path, uploaded_file.name)
                    if not os.path.exists(file_path):
                        with open(file_path, 'wb') as file:
                            file.write(uploaded_file.getbuffer())
                        st.success(f"{uploaded_file.name} uploaded successfully.")
                    else:
                        st.warning(f"{uploaded_file.name} already exists.")

    elif choose == "Download Document":
        subject_folders = [d for d in os.listdir(base_folder) if os.path.isdir(os.path.join(base_folder, d))]
        subject_folders.sort()
        subject_folders.insert(0, "Select Subject")
        
        selected_subject = st.selectbox("Select Subject", subject_folders, index=0, key='download_subject_selector')

        if selected_subject != "Select Subject":
            subject_folder_path = os.path.join(base_folder, selected_subject)
            pdf_files = [file for file in os.listdir(subject_folder_path) if file.endswith(".pdf")]
            pdf_files.insert(0, "Select Document")
            selected_file = st.selectbox("Select Document", pdf_files, index=0, key='download_selected_file')

            if selected_file != "Select Document":
                file_path = os.path.join(subject_folder_path, selected_file)
                with open(file_path, "rb") as file:
                    st.download_button(label="Download", data=file, file_name=selected_file, mime="application/pdf")
            else:
                st.info("Select a document to download.")

    elif choose == "View Documents":
        subject_folders = [d for d in os.listdir(base_folder) if os.path.isdir(os.path.join(base_folder, d))]
        subject_folders.sort()
        subject_folders.insert(0, "Select Subject")

        selected_subject = st.selectbox("Select Subject", subject_folders, index=0, key='view_subject_selector')

        if selected_subject != "Select Subject":
            subject_folder_path = os.path.join(base_folder, selected_subject)
            pdf_files = [file for file in os.listdir(subject_folder_path) if file.endswith(".pdf")]
            if pdf_files:
                st.write(f"Documents in the folder for {selected_subject}:")
                for pdf_file in pdf_files:
                    st.write(pdf_file)
            else:
                st.info("No documents found in the selected subject folder.")

    elif choose == "Delete Document":
        subject_folders = [d for d in os.listdir(base_folder) if os.path.isdir(os.path.join(base_folder, d))]
        subject_folders.sort()
        subject_folders.insert(0, "Select Subject")

        selected_subject = st.selectbox("Select Subject", subject_folders, index=0, key='delete_subject_selector')

        if selected_subject != "Select Subject":
            subject_folder_path = os.path.join(base_folder, selected_subject)
            pdf_files = [file for file in os.listdir(subject_folder_path) if file.endswith(".pdf")]
            pdf_files.insert(0, "Select Document")
            selected_file = st.selectbox("Select Document", pdf_files, index=0, key='delete_selected_file')

            if selected_file != "Select Document":
                os.remove(os.path.join(subject_folder_path, selected_file))
                st.success(f"{selected_file} has been successfully removed.")
            else:
                st.info("Select a document to delete.")
    if choose=="Add Word to Dictionary":
        if st.button("View Dictionary"):
            with open('dictionary.json','r') as f:
                    existing_dictionary = json.load(f)
            lowercase_dict = {key.lower(): value for key, value in existing_dictionary.items() if isinstance(value, str)}
            st.write(lowercase_dict)

        # Input field for searching a word
        search_word = st.text_input('Enter English Word to Search:', '')
        if search_word:
            with open('dictionary.json','r') as f:
                existing_dictionary = json.load(f)
            translation = existing_dictionary.get(search_word, 'Word not found')
            st.write(f"Hindi Translation for '{search_word}': {translation}")

        with open('dictionary.json','r') as f:
            st.session_state.existing_dictionary = json.load(f)

        # Function to save dictionary to a local file
        def save_dictionary_to_file(filename):
            #st.write(st.session_state.existing_dictionary)
            with open(filename, 'w') as f:
                json.dump(st.session_state.existing_dictionary, f)

        col1, col2 = st.columns(2)
        with col1:

            st.write("##### English Words (Separated by commas)")
            eng_txt = st.text_input(
                "Enter English Words",
                key="eng"
            )
        with col2:

            st.write("##### Hindi Words (Separated by commas)")
            punjabi_txt = st.text_input(
                "Enter Hindi Words",
                key="hi"
            )

        lowercase_dict = {key.lower(): value for key, value in st.session_state.existing_dictionary.items()}

        if st.button("Request to Add"):
        # Split the entered English and Punjabi words by commas and remove leading/trailing whitespace
            eng_words = [word.strip() for word in eng_txt.split(",")]
            punjabi_words = [word.strip() for word in punjabi_txt.split(",")]

            # Add each pair of words to the lowercase dictionary
            for eng_word, punjabi_word in zip(eng_words, punjabi_words):
                # Convert entered English word to lowercase
                eng_lower = eng_word.lower()
                # Add the entry to the lowercase dictionary
                st.session_state.existing_dictionary[eng_lower] = punjabi_word
            st.success("Words added to dictionary!")
            st.write("Dictionary:", st.session_state.existing_dictionary)
            save_dictionary_to_file("dictionary.json")
            st.success("Dictionary saved to local!")
                    
footer = """
    <style>
    body {
        margin: 0;
        padding-top: 70px;  /* Add padding to prevent content from being hidden behind the footer */
    }
    .footer {
        position: absolute;
        top: 80px;
        left: 0;
        width: 100%;
        background-color: #002F74;
        color: white;
        text-align: center;
        padding: 5px;
        font-weight: bold;
        z-index: 1000;  /* Ensure it is on top of other elements */
        display: flex;
        align-items: center;
        justify-content: space-between;
        flex-wrap: wrap;
    }
    .footer p {
        font-style: italic;
        font-size: 14px;
        margin: 0;
        flex: 1 1 50%;  /* Flex-grow, flex-shrink, flex-basis */
    }
    @media (max-width: 600px) {
        .footer p {
            flex-basis: 100%;
            text-align: center;
            padding-top: 10px;
        }
    }
    </style>
    <div class="footer">
        <p style="text-align: left;">Copyright © 2024 MPSeDC. All rights reserved.</p>
        <p style="text-align: right;">The responses provided on this website are AI-generated. User discretion is advised.</p>
    </div>
"""

st.markdown(footer, unsafe_allow_html=True)
